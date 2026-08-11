"""Extracts raw text from uploaded resume files. No FastAPI/request objects here —
pure functions operating on file paths, per the AI-module architecture rule."""

import docx
import fitz  # PyMuPDF
import pdfplumber


def extract_text_from_pdf(file_path: str) -> str:
    """Primary: PyMuPDF (fast, robust). Falls back to pdfplumber if PyMuPDF yields
    little or no text (e.g. some scanned/odd-encoded PDFs)."""
    text = ""
    try:
        with fitz.open(file_path) as doc:
            text = "\n".join(page.get_text() for page in doc)
    except Exception:
        text = ""

    if len(text.strip()) < 40:
        try:
            with pdfplumber.open(file_path) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception:
            pass
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    document = docx.Document(file_path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts).strip()


def extract_text(file_path: str, mime_type: str) -> str:
    if "pdf" in mime_type:
        return extract_text_from_pdf(file_path)
    if "word" in mime_type or file_path.lower().endswith(".docx"):
        return extract_text_from_docx(file_path)
    raise ValueError(f"Unsupported resume file type: {mime_type}")
